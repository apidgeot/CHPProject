# -*- coding: windows-1251 -*-

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import pandas as pd
from  backend import utils
import json
from functools import lru_cache
import requests
import time
from fastapi.responses import FileResponse
import os
from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from typing import Optional, Dict, Any, List
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon, FancyArrowPatch, Rectangle
import numpy as np
import textwrap
from matplotlib.lines import Line2D
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle
from datetime import datetime, timedelta
import matplotlib.font_manager as fm
import textwrap
import matplotlib.colors as mcolors
import numpy as np
from fastapi import FastAPI, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
import io
import traceback
import re
import aiohttp

app = FastAPI()

IMAGE_DIR = "generated_images"
os.makedirs(IMAGE_DIR, exist_ok=True)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["GET, POST"],
    allow_headers=["*"]
)


def get_all_questions(filename="C:\\Users\dnsuser\PycharmProjects\k�sarev\\backend\questions.json"):
    return json.load(open(filename, 'r', encoding='utf-8'))


questions_all = get_all_questions()

cache = {}


# ������ FastAPI
class Question(BaseModel):
    question: str
    answer_type: str
    options: Optional[dict[str, str]]
    next_question: Optional[str]
    time_related: Optional[bool]


class Answer(BaseModel):
    question: str
    answer_type: str
    next_question: Optional[str] = None
    options: Optional[Dict[str, Any]] = None  # ��� ���� �����������
    answer: Optional[str] = None
    answer_option: Optional[str] = None
    time_related: Optional[bool] = None


class Answers(BaseModel):
    questions: List[Answer]
    start_time: Optional[str] = None
    end_time: Optional[str] = None


class FormattedText(BaseModel):
    formatted_text: str


class OptionRequest(BaseModel):
    options: Dict[str, Dict[str, Any]]
    answer_text: str


class ReportData(BaseModel):
    template_path: str
    start_time: str
    finish_time: str
    questions_text: str
    list_about_questions: list[str]


async def llama_request_async(prompt, max_token=None):
    config = utils.read_config("C:\\Users\dnsuser\PycharmProjects\k�sarev\\backend\llama_config.json")
    url = config['llama_url']

    payload = {
        "model": config['llama_model'],
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "stream": False,
        "max_token": max_token
    }

    async with aiohttp.ClientSession() as session:
        async with session.post(url, json=payload) as response:
            result = await response.json()
            return result['message']['content']


async def llama_create_schedule_async(formatted_text):
    prompt = """
    �� ������ ����������� ��� �������� � ������� ������� �������� ���������� ��� ��������� ����� � ������� ��� �����������. ������� �� �� ��������� ��������, ��������� ������� � ���������� �������. ������ ������� ������ �������� �����, ������� ��� ��������, ������� ����� ���������� �� ������. ������� ����� � ������ ������� ���������� ������� �����
    ��� �����: 
    '''
    {result}
    '''
    """

    prompt = prompt.format(result=formatted_text)
    result = await llama_request_async(prompt)

    prompt = """
    �� ������ ���� �������� ������ ��� json ��������� ��� ��������� ����� �� �������:
    ������� ������ JSON-������:

    {{
        "process1": {{
            "10:00-12:00": "�������� ������� 1, ������������ � process1.", // ������ HH:MM ���������� �����
            "12:00-13:00": "�������� ������� 2, ������������ � process1."
        }},
        "process2": {{
            "09:00-11:00": "�������� ������� 1, ������������ � process2.",
            "11:00-13:00": "�������� ������� 2, ������������ � process2."
        }}
    }}

    ����������� ����������� ���� ������ ��� ����������� ��������� �����:\n
    '''
    {text}
    '''

    � ������ ����� ������ JSON ������ ����� ���� (������� �������):
    ```json
    ...
    ```
    \n
    � json ��������� ������� �������!
    """

    prompt = prompt.format(text=result)
    result = await llama_request_async(prompt)

    return result


# ������� ������ � ������
@lru_cache(maxsize=512)
def llama_request(prompt, max_token=None):
    # URL ������ �������
    config = utils.read_config("C:\\Users\dnsuser\PycharmProjects\k�sarev\\backend\llama_config.json")
    url = config['llama_url']

    payload = {
        "model": config['llama_model'],
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "stream": False,
        "max_token": max_token
    }

    # �������� POST-�������
    response = requests.post(
        url,
        data=json.dumps(payload)
    )

    print(response.json()['message']['content'])
    return response.json()['message']['content']


def llama_create_schedule(formatted_text):
    print('���� � llama_create_scedule: ', formatted_text)

    prompt = """
    �� ������ ����������� ��� �������� � ������� ������� �������� ���������� ��� ��������� ����� � ���� ������������. ������� ���������� �� ���������� ��������, ��������� ������� � ��������� �������. 

    ������ ������� ������ �������� �����, ������� ��� ��������, ��������� �������� � �������������. ����� ������ ����� ������ � ��������� ������� ������� � ������� HH:MM.

    �� ���������� �������� �� �������������� ������������ ��������, � ����������� �� ��������, ��������������� ��������� � �������������. ����� ��������� �������� ��� ����� �������� ���������� ���� ����� �� �������!

    �����:
    '''
    {result}
    '''
    """

    prompt = prompt.format(result=formatted_text)
    result = llama_request(prompt)

    prompt = """
    �� ������ ����������� �������� ������ JSON-��������� ��� ��������� �����, ��������������� ���������� �������:

    {{
        "process1": {{
            "10:00-11:00": "�������� ������� 1, ������������ � process1.",
            "11:00-12:00": "�������� ������� 2, ������������ � process1."
        }},
        "process2": {{
            "09:00-10:00": "�������� ������� 1, ������������ � process2.",
            "10:00-11:00": "�������� ������� 2, ������������ � process2."
        }}
    }}

    ����������:
    - ����� � ������� HH:MM-HH:MM.
    - ������� ������ ���� ������������� �� ���������� ���������.
    - ������� ������������� ������� ������� � JSON-������.
    - ����� ��������� �������� ��� ����� �������� ���������� ���� ����� �� �������!

    ��������� ��������� ������ ��� ����������� ��������� �����:
    '''
    {text}
    '''

    ����� ������ ���� � ������� JSON ������ ����� ����:
    ```json
    ...

    """

    prompt = prompt.format(text=result)
    result = llama_request(prompt)

    # �������� � ����������� ������� JSON
    try:
        json_data = json.loads(result)
    except json.decoder.JSONDecodeError:
        # ����������� ������� � �������������� ���������� ���������
        result = extract_json_object(result)
        result = re.sub(r"'", '"', result)  # �������������� ��������� ������� � �������
        #result = re.sub(r'(\d{2}:\d{2})', r'"\1"', result)  # ���������� ������� � �������
        json_data = json.loads(result)

    result = json.dumps(json_data)

    return result


def llama_create_isikava(formatted_text):
    prompt = """
    ��������� ���� �������� � ������� ����, �������� ������ ��� ��������� ��������-������������ ������ (��������� �������) � ��������� �������. ���������� �������� �������� � �������� ��������� �������� ������, ������ �� ������� �������� ����� ������, ���������� � �������

    ����������� ��������� ���� �������� � ������� ��� ����������� ����������� ��������:
    '''
    {text}
    '''
    """

    prompt = prompt.format(text=formatted_text)
    result = llama_request(prompt)

    prompt = """
    ��������� ������ ����, �������� ������ ��� ��������� ��������-������������ ������ (��������� �������) � ������� JSON. ���������� �������� �������� � �������� ��������� �������� ������, ������ �� ������� �������� ����� ������, ���������� � �������. ������ ����� ������ ������������ �������� ��������, � ������ ������� ������ ����� �������� ������ �������, ���������� � �������.

    ������� ������ JSON-������:

    {{
        "main_problem": "�������� ��������, ���������� � ��������.",
        "branshes": [
            {{
                "name": "branch_name_1",
                "reasons": [
                    "������� 1, ��������� � ������ ���������� ��������.",
                    "������� 2, ��������� � ������ ���������� ��������."
                ]
            }},
            {{
                "name": "branch_name_2",
                "reasons": [
                    "������� 1, ��������� � ������ ���������� ��������.",
                    "������� 2, ��������� � ������ ���������� ��������."
                ]
            }},
            {{
                "name": "branch_name_3",
                "reasons": [
                    "������� 1, ��������� � ������� ���������� ��������.",
                    "������� 2, ��������� � ������� ���������� ��������."
                ]
            }}
        ]
    }}

    ������ ������ ���� ������ json!

    ������ ����� ����� ��������� �������, ���������������� �� ����������, ����� ��� "������������ ������", "���� ����������", "���������� � ������������" � �.�.
    ����������� ��������� ������ ��� ����������� ���������:
    '''
    {text}
    '''

    � ������ ����� ������ JSON ������ ����� ���� (������� �������):
    ```json
    ...
    ```
    � json ��������� ������� �������!
    """

    prompt = prompt.format(text=result)
    result = llama_request(prompt)

    # �������� � ����������� ������� JSON
    try:
        json_data = json.loads(result)
    except json.decoder.JSONDecodeError:
        # ����������� ������� � �������������� ���������� ���������
        result = extract_json_object(result)
        result = re.sub(r"'", '"', result)  # �������������� ��������� ������� � �������
        #result = re.sub(r'(\d{2}:\d{2})', r'"\1"', result)  # ���������� ������� � �������
        json_data = json.loads(result)


    result = json.dumps(json_data)

    return result


# ������� ��� ������������
def save_isikava_image(data, filename):
    draw_isikava(data)
    image_path = os.path.join(IMAGE_DIR, filename)
    plt.savefig(image_path, bbox_inches='tight')
    plt.close()
    return image_path


def save_timeline_image(data, filename):
    create_timeline(data)
    image_path = os.path.join(IMAGE_DIR, filename)
    plt.savefig(image_path, bbox_inches='tight')
    plt.close()
    return image_path


def draw_isikava(data):
    # ��������� �������� ��������
    main_problem = data['main_problem']

    # ��������� ����� � �������
    branches = data['branches']
    num_branches = len(branches)

    # ������� ������������ ���������� ������
    max_causes = max(len(branch['reasons']) for branch in branches)

    # ����������� ������ ������ � ����������� �� ������
    fig_width = 14
    fig_height = max(6, num_branches * 2)
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis('off')

    # ��������� ���
    ax.set_facecolor('#f0f8ff')

    # ������ �������� ������ � ����������
    spine_length = 20
    ax.plot([0, spine_length], [0, 0], color='#1f4e79', linewidth=4, solid_capstyle='round')

    # ������� �������� �������� � �������� �����
    wrapped_main_problem = textwrap.fill(main_problem, width=30)
    bbox_props = dict(boxstyle="round,pad=1", fc="#add8e6", ec="#1f4e79", lw=2)
    ax.text(spine_length, 0, wrapped_main_problem, fontsize=10, fontweight='bold',
            va='center', ha='left', bbox=bbox_props, color='#1f4e79')

    # ������ ����� ���� � ����� ��������� ��������
    tail = Polygon([[-1, -0.5], [0, 0], [-1, 0.5]], closed=True, fc='#1f4e79', ec='#1f4e79')
    ax.add_patch(tail)

    # �������� ������� ��� ������
    colors = ['#87ceeb', '#4682b4', '#5f9ea0', '#b0c4de', '#add8e6']

    # ������� ������ ����� ������
    branch_x_positions = np.linspace(spine_length * 0.1, spine_length * 0.9, num_branches)

    for i, branch in enumerate(branches):
        branch_name = branch['name']
        causes = branch['reasons']
        branch_x = branch_x_positions[i]
        branch_color = colors[i % len(colors)]  # ��������� ���� � �����

        # ����������, �������� �� ����� ����� ��� ����
        if i % 2 == 0:
            branch_y = 4 + len(causes)  # ������ ������������ ��� ������
            y_direction = 1
            valign = 'bottom'
        else:
            branch_y = -4 - len(causes)
            y_direction = -1
            valign = 'top'

        # ������ ����� ����� � ����� ������� ������ � ��������
        arrow = FancyArrowPatch((branch_x, 0), (branch_x, branch_y), arrowstyle='->',
                                mutation_scale=20, linewidth=2, color=branch_color)
        ax.add_patch(arrow)

        # ��������� �������� ����� � �����
        wrapped_branch_name = textwrap.fill(branch_name, width=20)
        bbox_branch = dict(boxstyle="round,pad=0.5", fc="#f0f8ff", ec=branch_color, lw=2)
        ax.text(branch_x, branch_y + y_direction * 0.5, wrapped_branch_name, fontsize=12,
                va=valign, ha='center', fontweight='bold', color=branch_color, bbox=bbox_branch)

        # ������ ������� ����� ����� � ����� �������� �����������
        num_causes = len(causes)
        cause_positions = np.linspace(0, branch_y, num_causes + 2)[1:-1]  # ��������� ��������� � �������� �����

        for j, cause in enumerate(causes):
            cause_x = branch_x
            cause_y = cause_positions[j]
            # ������ ����� �������, ������������� �����
            cause_line_length = 3  # ����� ����� �������
            angle = np.degrees(np.arctan2(cause_y, cause_line_length))
            if y_direction > 0:
                end_x = cause_x - cause_line_length * np.cos(np.radians(15))
                end_y = cause_y - cause_line_length * np.sin(np.radians(15))
            else:
                end_x = cause_x - cause_line_length * np.cos(np.radians(15))
                end_y = cause_y + cause_line_length * np.sin(np.radians(15))

            ax.plot([cause_x, end_x], [cause_y, end_y], color=branch_color, linewidth=2)

            # ��������� ����� ������� � �����
            wrapped_cause = textwrap.fill(cause, width=30)
            bbox_props_cause = dict(boxstyle="round,pad=0.3", fc="#fffacd", ec=branch_color, lw=1.5)
            ax.text(end_x - 0.5, end_y, wrapped_cause, fontsize=7, va='center',
                    ha='right', bbox=bbox_props_cause, color='#1f4e79')

    # ����������� ������� ����
    y_max = max(5, abs(branch_y)) + max_causes * 0.5
    ax.set_xlim(-5, spine_length + 10)
    ax.set_ylim(-y_max, y_max)

    plt.tight_layout()


def create_timeline(data):
    processes = list(data.keys())
    num_processes = len(processes)
    
    # �������� ��� ��������� ����� � ���������� ����� ��������� ��������
    all_times = set()
    for events in data.values():
        for time_interval in events.keys():
            start_str, end_str = [t.strip() for t in time_interval.split('-')]
            start_time = datetime.strptime(start_str, '%H:%M')
            end_time = datetime.strptime(end_str, '%H:%M')
            all_times.update([start_time, end_time])
    all_times = sorted(all_times)
    
    # ���������� ����������� � ������������ �����
    min_time = min(all_times)
    max_time = max(all_times)
    total_minutes = (max_time - min_time).total_seconds() / 60
    
    # ������� ������ � ��� � ������������ ��������� � ����� ������� �������������
    fig, ax = plt.subplots(figsize=(10, num_processes * 1.5))
    ax.set_title('��������� ����� ���������', fontsize=14, fontweight='bold', color='#1f4e79', pad=20)
    
    # ����������� ������� � ��������� ��� ��� ����������
    ax.set_xlim(0, total_minutes)
    ax.set_ylim(0.5, num_processes + 0.5)
    ax.axis('off')
    
    # ��������� ��� ��� �������
    ax.set_facecolor('#f0f8ff')
    
    # ������ �������������� ����� ��� ��������� � ����������� �����������
    for i in range(1, num_processes + 1):
        ax.hlines(i, 0, total_minutes, colors='#d3d3d3', linestyles='dotted', linewidth=1)
        ax.text(-2, i, processes[i - 1], verticalalignment='center', horizontalalignment='right', fontsize=10, color='#1f4e79')
    
    # ��������� ����� ������� ������ 30 �����
    time_labels = []
    time_positions = []
    current_time = min_time.replace(minute=(min_time.minute // 30) * 30, second=0, microsecond=0)
    while current_time <= max_time:
        pos = (current_time - min_time).total_seconds() / 60
        time_positions.append(pos)
        time_labels.append(current_time.strftime('%H:%M'))
        current_time += timedelta(minutes=30)
    
    # ������ ������������ ����� ��� ����� �������
    for pos in time_positions:
        ax.vlines(pos, 0.5, num_processes + 0.5, colors='#d3d3d3', linestyles='dotted', linewidth=1)
    
    # ��������� ����� ������� ����� �������
    for pos, label in zip(time_positions, time_labels):
        ax.text(pos, 0.4, label, verticalalignment='top', horizontalalignment='center', fontsize=8, color='#1f4e79')
    
    # �������� ������� ��� �������
    event_colors = ['#add8e6', '#87ceeb', '#4682b4', '#b0c4de', '#5f9ea0']
    
    # ������������ ������� ��� ������� ��������
    for i, process in enumerate(processes, 1):
        events = data[process]
        event_list = []
        for time_interval, description in events.items():
            start_str, end_str = [t.strip() for t in time_interval.split('-')]
            start_time_event = datetime.strptime(start_str, '%H:%M')
            end_time_event = datetime.strptime(end_str, '%H:%M')
            start_pos = (start_time_event - min_time).total_seconds() / 60
            end_pos = (end_time_event - min_time).total_seconds() / 60
            width = end_pos - start_pos
            event_list.append({'start': start_pos, 'end': end_pos, 'width': width, 'description': description})
        
        # ��������� ������� �� ������� ������
        event_list.sort(key=lambda x: x['start'])
        
        # ������������ �������������� �������
        layers = []
        for event in event_list:
            placed = False
            for layer_index, layer in enumerate(layers):
                # ���������, ������������ �� ������� � ������� � ���� ����
                if all(event['start'] >= e['end'] or event['end'] <= e['start'] for e in layer):
                    layer.append(event)
                    event['layer'] = layer_index
                    placed = True
                    break
            if not placed:
                # ������� ����� ����
                layers.append([event])
                event['layer'] = len(layers) - 1
        
        # ������������ ������� � ������ �����
        max_layer = max(event['layer'] for event in event_list) if event_list else 0
        layer_height = 0.6 / (max_layer + 1)
        
        for event in event_list:
            layer_offset = layer_height * event['layer']
            rect_y = i - 0.3 + layer_offset
            rect_height = layer_height * 0.9  # ������� ��������� ������ ��� ��������
            # �������� ���� ��� �������
            color_index = event['layer'] % len(event_colors)
            event_color = event_colors[color_index]
            
            # ��������� ����
            shadow = Rectangle(
                (event['start'] + 0.5, rect_y - 0.05),
                event['width'],
                rect_height,
                facecolor='gray',
                edgecolor=None,
                linewidth=0,
                zorder=1,
                alpha=0.2
            )
            ax.add_patch(shadow)
            
            # ������ ������������� �������
            rect = Rectangle(
                (event['start'], rect_y),
                event['width'],
                rect_height,
                facecolor=event_color,
                edgecolor='#1f4e79',
                linewidth=1,
                zorder=2,
                alpha=0.85,
                clip_on=False,
                joinstyle='round',
                capstyle='round'
            )
            ax.add_patch(rect)
            
            # ��������� ����� � ���������
            max_chars = max(5, int(event['width'] / (total_minutes / 50)) * 2)  # ��������� ������ ��� ��������
            wrapped_text = '\n'.join(textwrap.wrap(event['description'], width=max_chars))
            
            ax.text(
                event['start'] + event['width'] / 2,
                rect_y + rect_height / 2,
                wrapped_text,
                verticalalignment='center',
                horizontalalignment='center',
                fontsize=8,
                color='#1f4e79',
                wrap=True,
                clip_on=False
            )
    
    # ��������� �������
    #custom_patches = [Rectangle((0,0),1,1, facecolor=color, edgecolor='#1f4e79') for color in event_colors]
    #legend_labels = ['���� {}'.format(i+1) for i in range(len(event_colors))]
    #ax.legend(custom_patches, legend_labels, loc='upper right', fontsize=8, title='���� �������', title_fontsize=9)
    
    plt.tight_layout()
    #plt.show()


# ���������� ������ �� ��������
@app.post("/api/fill_report")
async def fill_report(data: ReportData):
    # ��������� ������ ���������
    doc = Document(data.template_path)

    # ������������ �������
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                # ���� ������ � ������� "����� ����� �" � ��������� ����� ������ � ��������� ������
                if '����� ����� �' in cell.text and i + 1 < len(row.cells):
                    row.cells[i + 1].text = data.start_time.split(':')[0]
                    row.cells[i + 3].text = data.start_time.split(':')[1]
                # ���� ������ � ������� "����� ������� �" � ��������� ����� ��������� � ��������� ������
                if '����� ������� �' in cell.text and i + 1 < len(row.cells):
                    row.cells[i + 1].text = data.finish_time.split(':')[0]
                    row.cells[i + 3].text = data.finish_time.split(':')[1]

    # ������ ������ ��������
    for para in doc.paragraphs:
        if '���� �������� ���������:' in para.text:
            para.add_run(f'\n\n{data.questions_text}')

    # ������ ������ � ��������
    about_fields = [
        '1. �������, ���, �������� (��� �������)',
        '2. ���� ��������',
        '3. ����� ��������',
        '4. ����� ���������� � (���) �����������',
        '5. �������',
        '6. ����������� �����',
        '7. �����������',
        '8. �����������',
        '9. �������� ���������',
        '10. ����� ����� ��� ������',
        '11. ���������, ���������'
    ]

    idx = 0
    for para in doc.paragraphs:
        if idx < len(data.list_about_questions):
            for field in about_fields:
                if field in para.text:
                    para.text = para.text.replace(field, f'{field}: {data.list_about_questions[idx]}')
                    idx += 1
                    if idx >= len(data.list_about_questions):
                        break

    # ��������� �������� �� ��������� �����
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # ���������� ��������� �� ������ �����

    # ���������� ���� ��� ����������
    headers = {
        'Content-Disposition': 'attachment; filename="filled_report.docx"'
    }

    return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


# ��������� json-�������� �� �������
@app.post('/api/create-schedule')
def create_schedule(text: FormattedText):
    json_data = llama_create_schedule(text.formatted_text)
    return {'timeline': json_data}


@app.post('/api/create-isikava')
def create_isikava(text: FormattedText):
    json_data = llama_create_isikava(text.formatted_text)
    return {'isikava': json_data}


def extract_json_object(s):
    # ������� ��� ������� ����� ������ "{" � ��������� "}"
    match = re.search(r'\{.*\}', s, re.DOTALL)
    if match:
        return match.group(0)
    return None


# ��������� �����������
@app.post('/api/generate-ishikawa-image')
def generate_ishikawa_image(text: FormattedText):
    try:
        json_data = extract_json_object(text.formatted_text)
        data = json.loads(json_data)

        # Save the image
        image_path = save_isikava_image(data, "ishikawa_diagram.png")
        
        return FileResponse(image_path, media_type="image/png", filename="ishikawa_diagram.png")
    except Exception as e:
        print(str(e))
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post('/api/generate-timeline-image')
def generate_timeline_image(text: FormattedText):
    try:
        #json_data = llama_create_schedule(text.formatted_text)
        json_data = extract_json_object(text.formatted_text)
        data = json.loads(json_data)

        # Save the image
        image_path = save_timeline_image(data, "timeline.png")
        
        return FileResponse(image_path, media_type="image/png", filename="timeline.png")
    except Exception as e:
        print(str(e))
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# Endpoint to choose an option from list with llama model launched locally
@app.post("/api/choose-option")
def choose_option(request: OptionRequest):
    print('here')
    options = request.options
    answer = request.answer_text

    prompt = "������ ��� ������� � ������ �� ��������� ����� ����� ����� ������������ ����������!\n\n"
    prompt += f"����� ������������: {answer}\n\n"

    prompt += "�������� ������:\n"
    for option in options:
        prompt += f"{option} - {options[option]['title']}\n"
     
    prompt += "\n����� ��� ������ �����/����� ���������� ����� ��������. ���� �� �� ������ ������� ����� �� � ������ �� ������������ �������� (����� ��������� ������������), �� ���������� ������ ������ '0'"

    # return 0

    max_attempts = 3
    attempts = 0
    while attempts < 3:
        ans = llama_request(prompt, 10)
        print(ans)
        try:
            int(ans)
            return int(ans)
        except Exception:
            attempts += 1
    return 0


# Endpoint to get the question
# Question id example: branch-sheet-block-question_num
@app.get("/api/get-question/{question_id}")
def get_questions(question_id: str):
    branch = question_id.split('-')[0]
    sheet = question_id.split('-')[1]
    block = question_id.split('-')[2]
    question_num = question_id.split('-')[3]

    print(branch, sheet, block, question_num)

    return questions_all[branch][sheet][block][question_num]


# FastAPI ������� ��� ������ � ���������� �������
@app.post("/api/submit-answers")
def submit_answers(answers: Answers):
    try:
        print("Saving answers")
        # ������ ������������� ������������ � ������� Pydantic
        validated_answers = answers.dict()

        formatted_text = ""
        for x in validated_answers['questions']:
            formatted_text += f"������: {x['question']}\n"
            if x["answer_type"] == "text":
                formatted_text += f"�����: {x['answer']}\n\n"
            elif x["answer_type"] in ['test', 'test-rec']:
                formatted_text += f"�����: {x['options'][x['answer_option']]['title']}\n\n"

        # ��������� ������ � JSON ���� (�����������)
        with open(f'{str(validated_answers["questions"][0]["answer"])}_answers.json', 'w', encoding='utf-8') as f:
            json.dump(validated_answers, f, ensure_ascii=False, indent=4)

        # ��������� ������ � ���� ������
        db = SessionLocal()
        new_log_entry = AnswerLog(
            answers_text=formatted_text,
            start_time=validated_answers["start_time"],
            end_time=validated_answers["end_time"]
        )
        db.add(new_log_entry)
        db.commit()
        db.refresh(new_log_entry)
        db.close()

        return {"message": "Answers received and saved successfully."}

    except Exception as e:
        # �������� ���������� � ����� ������
        raise HTTPException(status_code=422, detail=f"������ ��� ��������� ������: {str(e)}")


# �������������� ������� ������������
@app.post('/api/format-answers')
def format_answers(answers: Answers):
    try:
        print("Formatting answers")

        # ������ ������������� ������������ � ������� Pydantic
        validated_answers = answers.dict()

        formatted_text = ""

        for x in validated_answers['questions']:
            formatted_text += f"������: {x['question']}\n"

            if x["answer_type"] == "text":
                formatted_text += f"�����: {x['answer']}\n\n"
            elif x["answer_type"] in ['test', 'test-rec']:
                formatted_text += f"�����: {x['options'][x['answer_option']]['title']}\n\n"

        return {"formatted_text": formatted_text}
    
    except Exception as e:
        # �������� ���������� � ����� ������
        raise HTTPException(status_code=422, detail=f"������ ��� ��������� ������: {str(e)}")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("app:app", host="127.0.0.1", port=8000, reload=True)